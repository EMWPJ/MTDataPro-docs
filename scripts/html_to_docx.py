#!/usr/bin/env python3
"""Convert Sphinx HTML docs to DOCX."""

import os
import re
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
HTML_DIR = os.path.join(BASE_DIR, "..", "_build", "html")
OUTPUT = os.path.join(BASE_DIR, "..", "_build", "docx", "mtdatapro.docx")

CHAPTER_FILES = [
    ("chapters/chapter1.html", "第1章 软件简介"),
    ("chapters/chapter2.html", "第2章 MT原理"),
    ("chapters/chapter3.html", "第3章 数据处理"),
    ("chapters/chapter4.html", "第4章 附录"),
    ("chapters/chapter5.html", "第5章 RMT数据处理"),
    ("chapters/chapter6.html", "第6章 数据导出"),
    ("chapters/chapter7.html", "第7章 数据处理流程"),
    ("chapters/chapter8.html", "第8章 视图功能"),
]


class HTMLContentExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.in_article = False
        self.in_section = False
        self.skip_tags = {"script", "style", "nav", "header", "footer", "aside"}
        self.current_skip = 0
        self.content = []
        self.current_tag_stack = []
        self.text_buffer = ""
        self.in_code = False

    def handle_starttag(self, tag, attrs):
        if tag in self.skip_tags:
            self.current_skip += 1
        self.current_tag_stack.append(tag)
        if tag == "code" or tag == "pre":
            self.in_code = True

    def handle_endtag(self, tag):
        if tag in self.skip_tags and self.current_skip > 0:
            self.current_skip -= 1
        if self.current_tag_stack and self.current_tag_stack[-1] == tag:
            self.current_tag_stack.pop()
        if tag == "code" or tag == "pre":
            self.in_code = False
        if tag in ("p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "td", "th"):
            if self.text_buffer.strip():
                self.content.append(("para", self.text_buffer.strip()))
                self.text_buffer = ""
        if tag == "br":
            self.text_buffer += " "
        if tag == "tr":
            self.content.append(("para", ""))

    def handle_data(self, data):
        if self.current_skip == 0:
            self.text_buffer += data

    def handle_entityref(self, name):
        entities = {
            "nbsp": " ",
            "amp": "&",
            "lt": "<",
            "gt": ">",
            "quot": '"',
            "apos": "'",
            "ndash": "–",
            "mdash": "—",
            "hellip": "…",
            "lsquo": """, 'rsquo': """,
            "ldquo": '"',
            "rdquo": '"',
            "copy": "©",
            "reg": "®",
            "times": "×",
            "divide": "÷",
            "plusmn": "±",
            "sum": "∑",
            "int": "∫",
            "partial": "∂",
        }
        self.text_buffer += entities.get(name, f"&{name};")

    def get_content(self):
        return self.content


def clean_text(text):
    """Clean HTML entities and extra whitespace."""
    text = re.sub(
        r"&[a-zA-Z]+;",
        lambda m: {
            "nbsp": " ",
            "amp": "&",
            "lt": "<",
            "gt": ">",
            "quot": '"',
            "ndash": "–",
            "mdash": "—",
            "hellip": "…",
            "times": "×",
            "divide": "÷",
            "plusmn": "±",
            "sum": "∑",
            "int": "∫",
        }.get(m.group()[1:-1], m.group()),
        text,
    )
    text = re.sub(r"&#(\d+);", lambda m: chr(int(m.group(1))), text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract_title(html_content):
    """Extract page title from HTML."""
    match = re.search(r"<title>([^<]+)</title>", html_content)
    if match:
        return match.group(1).split(" — ")[0]
    return None


def add_heading(doc, text, level=1):
    p = doc.add_heading(clean_text(text), level=level)
    return p


def add_paragraph(doc, text, bold=False):
    text = clean_text(text)
    if not text:
        return None
    p = doc.add_paragraph()
    if bold:
        run = p.add_run(text)
        run.bold = True
    else:
        p.add_run(text)
    return p


def process_chapter(doc, filepath):
    """Process a chapter HTML file and add content to doc."""
    full_path = os.path.join(HTML_DIR, filepath)
    if not os.path.exists(full_path):
        print(f"  [SKIP] {filepath} not found")
        return

    with open(full_path, "r", encoding="utf-8") as f:
        html = f.read()

    title = extract_title(html)
    if title:
        add_heading(doc, title, 1)

    # Extract article body
    body_match = re.search(
        r'<div[^>]*itemprop="articleBody"[^>]*>(.*?)</div>\s*</article>',
        html,
        re.DOTALL,
    )
    if not body_match:
        body_match = re.search(r"<body[^>]*>(.*?)</body>", html, re.DOTALL)
    if not body_match:
        print(f"  [SKIP] No body found in {filepath}")
        return

    body = body_match.group(1)

    # Remove script and style tags
    body = re.sub(r"<script[^>]*>.*?</script>", "", body, flags=re.DOTALL)
    body = re.sub(r"<style[^>]*>.*?</style>", "", body, flags=re.DOTALL)

    # Process headings
    for h_level in range(1, 5):
        pattern = rf"<h{h_level}[^>]*>(.*?)</h{h_level}>"
        matches = re.finditer(pattern, body, re.DOTALL)
        for m in matches:
            heading_text = re.sub(r"<[^>]+>", "", m.group(1))
            heading_text = clean_text(heading_text)
            if heading_text:
                add_heading(doc, heading_text, h_level)
                body = (
                    body[: m.start()] + "<!--HEADING_PLACEHOLDER-->" + body[m.end() :]
                )

    # Process paragraphs and other elements
    blocks = re.split(r"(?=<(?:p|h[1-6]|pre|table|ul|ol|blockquote|div)[^>]*>)", body)

    for block in blocks:
        block = block.strip()
        if not block or "<!--HEADING_PLACEHOLDER-->" in block:
            continue

        if block.startswith("<p"):
            text = re.sub(r"<[^>]+>", "", block)
            text = clean_text(text)
            if text:
                # Check if it's a caption or label
                if re.match(r"^图\s*\d+", text) or re.match(r"^表\s*\d+", text):
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.italic = True
                    run.font.size = Pt(9)
                else:
                    add_paragraph(doc, text)
        elif block.startswith("<pre"):
            text = re.sub(r"<[^>]+>", "", block)
            text = clean_text(text)
            if text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
        elif block.startswith("<table"):
            rows = re.findall(r"<tr[^>]*>(.*?)</tr>", block, re.DOTALL)
            for row in rows:
                cells = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", row, re.DOTALL)
                if cells:
                    cell_texts = [clean_text(re.sub(r"<[^>]+>", "", c)) for c in cells]
                    p = doc.add_paragraph()
                    p.add_run(" | ".join(cell_texts))
        elif block.startswith(("<ul", "<ol")):
            items = re.findall(r"<li[^>]*>(.*?)</li>", block, re.DOTALL)
            for item in items:
                text = clean_text(re.sub(r"<[^>]+>", "", item))
                if text:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(text)
        elif block.startswith("<blockquote"):
            text = re.sub(r"<[^>]+>", "", block)
            text = clean_text(text)
            if text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.italic = True

    doc.add_page_break()


def main():
    print("Creating DOCX from Sphinx HTML docs...")
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

    doc = Document()
    doc.core_properties.title = "MTDataPro 中文手册"
    doc.core_properties.author = "王培杰"
    doc.core_properties.subject = "MTDataPro 大地电磁数据处理软件"

    # Title page
    title_p = doc.add_heading("MTDataPro 中文手册", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("版本 1.9.5")
    doc.add_paragraph("© 版权所有 2026, 王培杰")
    doc.add_page_break()

    for filepath, title in CHAPTER_FILES:
        print(f"Processing {filepath}...")
        process_chapter(doc, filepath)

    doc.save(OUTPUT)
    print(f"Saved to {OUTPUT}")


if __name__ == "__main__":
    main()
